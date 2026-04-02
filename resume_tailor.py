"""
resume_tailor.py

Takes the base resume DOCX and a company profile, uses the Anthropic API to make
minimal targeted adjustments (skills reordering, summary tweak, 1-2 bullet
rephrases) without fabricating anything, then validates the result before saving.

Paths are configurable via environment variables:
  RESUME_DOCX_PATH  - path to the base resume .docx file
  RESUME_OUTPUT_DIR - directory to write tailored resumes (default: ./resume_outputs)

Usage:
    from resume_tailor import tailor_resume
    path = tailor_resume("Trail of Bits", research_dict, role_type="offensive security")
"""

from __future__ import annotations

import json
import logging
import os
import re
import shutil
import tempfile
from pathlib import Path

import anthropic
from docx import Document

logger = logging.getLogger(__name__)

# Paths — configurable via env vars
BASE_DOCX = Path(os.getenv("RESUME_DOCX_PATH", "resume/Francis_Konikkara_Resume.docx"))
OUTPUT_DIR = Path(os.getenv("RESUME_OUTPUT_DIR", "resume_outputs"))

# Validation thresholds
MAX_WORD_DELTA_PCT = 0.08
MAX_BULLET_DELTA = 2
REQUIRED_SECTIONS = {"SKILLS", "EXPERIENCE", "PROJECTS", "EDUCATION"}


# ---------------------------------------------------------------------------
# Extraction helpers
# ---------------------------------------------------------------------------

def _paragraphs_text(doc: Document) -> str:
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def _bullet_count(doc: Document) -> int:
    return sum(
        1 for p in doc.paragraphs
        if (p.style.name or "") == "List Paragraph" and p.text.strip()
    )


def _section_headers(doc: Document) -> set[str]:
    return {
        p.text.strip().upper()
        for p in doc.paragraphs
        if (p.style.name or "") in ("Heading 1", "Heading 2") and p.text.strip()
    }


def _word_count(doc: Document) -> int:
    return len(_paragraphs_text(doc).split())


def _extract_skills_paragraph(doc: Document) -> tuple[int, str]:
    """Return (paragraph_index, text) of the Technical Skills line."""
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text.startswith("Technical Skills") or text.startswith("Skills"):
            return i, text
    return -1, ""


def _extract_summary_paragraph(doc: Document) -> tuple[int, str]:
    """Return (paragraph_index, text) of the professional summary if present."""
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if len(text) > 80 and (p.style.name or "") not in ("Heading 1", "Heading 2", "List Paragraph"):
            if any(kw in text.lower() for kw in ["security", "engineer", "research", "vulnerability", "penetration"]):
                return i, text
    return -1, ""


# ---------------------------------------------------------------------------
# Anthropic-powered tailoring
# ---------------------------------------------------------------------------

ABOUT_FRANCIS = """
Name: Francis Anthony Konikkara
Current: MEng Cybersecurity Engineering, University of Maryland (graduating May 2027)
CPT-authorized for summer 2026, F-1 visa, no sponsorship needed.

Experience:
- Analyst, Deloitte Touche Tomatsu LLP — Oct 2023 to Jul 2025
  Supported security-adjacent analytics and ML workflows using Python and JavaScript.
  Worked across AWS and Azure for enterprise-scale analytics. Secure data handling,
  model evaluation, and feedback pipelines for ML systems.
- Full Stack Development Intern, Zedex Info Pvt. Ltd — Jul 2022 to Apr 2023
  Secure web-based ERP using React, Node.js, GraphQL. Authentication, authorization,
  API security controls, secure logging. Third-party API integration with validation.
- Information Security Analyst Intern, CybersmithSecure Pvt Ltd — Apr 2021 to Aug 2021
  Vulnerability assessments and penetration testing using Nessus, Burp Suite, Metasploit.
  Technical reports, threat categorization with senior analysts.

Projects:
- DFIR Automation Framework: Full DFIR automation. Modules: disk/memory/network forensics,
  IR workflow (PICERL), malware analysis, threat hunting (7 hypotheses), Wazuh/TheHive/Shuffle
  SOAR integration, VirusTotal/MISP/AlienVault OTX threat intel.
- Secure CI/CD Pipeline (DevSecOps): Compliance-ready CI/CD with SAST, secrets scanning,
  container scanning. Policy-as-code gates aligned with SOC 2 and ISO/IEC 27001.
- SIEM & Detection Engineering Lab: ELK-based SIEM, MITRE ATT&CK detection rules,
  attack simulation, SOC triage and IR workflows.
- Malware Analysis Lab: Static and dynamic analysis of Windows/Linux malware, IOC extraction,
  reverse engineering and sandboxing.

Skills: Python, Go, JavaScript, TypeScript, Bash, PowerShell, SQL, ReactJS, Django, Flask, Node.js |
        Nessus, Metasploit, Burp Suite, Wireshark, Nmap, OSINT |
        Splunk, ELK, Wazuh, TheHive, Shuffle | AWS, Azure, GCP, Docker, Kubernetes, Terraform |
        MITRE ATT&CK, OWASP Top 10, Sigma, YARA, Volatility 3
Certs: CEH, Practical Ethical Hacking (TCM Security)
CTF: UMBC Nightwing CTF (3rd Place), Fword CTF, Cyber Apocalypse CTF
"""


def _call_claude(prompt: str) -> str:
    """Call Anthropic API with a prompt and return the text response."""
    try:
        client = anthropic.Anthropic()  # reads ANTHROPIC_API_KEY from env
        message = client.messages.create(
            model="claude-3-5-sonnet-20241022",
            max_tokens=2048,
            messages=[{"role": "user", "content": prompt}],
        )
        return message.content[0].text
    except Exception as e:
        logger.error(f"Anthropic API call failed: {e}")
        return ""


def _call_claude_for_tailoring(
    company_name: str,
    company_research: dict,
    role_type: str,
    current_skills_line: str,
    current_summary: str,
) -> dict:
    """
    Ask Claude for minimal targeted resume changes. Returns a dict with:
      - skills_line: rewritten skills line (same facts, reordered/rephrased)
      - summary: rewritten summary (same facts, slightly reframed for this company)
      - bullet_tweaks: list of {original: str, revised: str} — max 2 bullets
    """
    company_context = f"""
Company: {company_name}
What they do: {company_research.get('description', '')[:600]}
Recent news: {company_research.get('recent_news', '')[:300]}
Open roles: {company_research.get('open_roles', '')[:300]}
Target role type: {role_type}
"""

    prompt = f"""You are tailoring Francis Konikkara's resume for a specific company. Make MINIMAL changes — reorder and rephrase only, never invent new facts.

Rules:
- Do NOT add any experience, skills, tools, companies, or metrics that aren't already in the resume
- Do NOT change any dates, numbers, or proper nouns
- Only reorder skills within the existing list to put the most relevant first
- Only rephrase existing bullet points to emphasize the aspect most relevant to this company (same facts, different framing)
- Max 2 bullet tweaks total
- If the resume has a professional summary, reframe it for this company's focus area — same facts, different emphasis
- Keep everything at exactly the same length or shorter

About Francis (source of truth — nothing outside this may appear in the resume):
{ABOUT_FRANCIS}

Current skills line:
{current_skills_line}

Current professional summary (may be empty):
{current_summary}

Company context:
{company_context}

Return ONLY valid JSON in this exact format, nothing else:
{{
  "skills_line": "Technical Skills: [reordered/same skills]",
  "summary": "[reframed summary or empty string if no summary exists]",
  "bullet_tweaks": [
    {{"original": "exact text of bullet to change", "revised": "revised text — same facts, different framing"}}
  ]
}}"""

    try:
        raw = _call_claude(prompt)
        match = re.search(r"\{.*\}", raw, re.DOTALL)
        if match:
            return json.loads(match.group(0))
    except Exception as e:
        logger.error(f"Claude tailoring call failed: {e}")
    return {"skills_line": current_skills_line, "summary": current_summary, "bullet_tweaks": []}


# ---------------------------------------------------------------------------
# DOCX modification
# ---------------------------------------------------------------------------

def _replace_paragraph_text(para, new_text: str) -> None:
    """Replace paragraph text while preserving runs' formatting."""
    if not para.runs:
        para.text = new_text
        return
    first_run = para.runs[0]
    first_run.text = new_text
    for run in para.runs[1:]:
        run.text = ""


def _apply_tailoring(doc: Document, changes: dict) -> Document:
    """Apply Claude's suggested changes to the document."""
    skills_line = changes.get("skills_line", "")
    summary = changes.get("summary", "")
    bullet_tweaks = changes.get("bullet_tweaks", [])

    tweak_map = {
        re.sub(r"\s+", " ", t["original"].strip()): t["revised"]
        for t in bullet_tweaks
        if t.get("original") and t.get("revised")
    }

    for para in doc.paragraphs:
        text = para.text.strip()
        norm = re.sub(r"\s+", " ", text)

        if (text.startswith("Technical Skills") or text.startswith("Skills")) and skills_line:
            _replace_paragraph_text(para, skills_line)
            continue

        if summary and text == re.sub(r"\s+", " ", summary.strip()):
            _replace_paragraph_text(para, summary)
            continue

        for orig_norm, revised in tweak_map.items():
            if norm[:60] == orig_norm[:60] or orig_norm[:60] in norm:
                _replace_paragraph_text(para, revised)
                break

    return doc


# ---------------------------------------------------------------------------
# Validation
# ---------------------------------------------------------------------------

def _validate(original: Document, tailored: Document) -> list[str]:
    """Return list of validation errors. Empty = pass."""
    errors = []

    orig_sections = _section_headers(original)
    tail_sections = _section_headers(tailored)
    missing = orig_sections - tail_sections
    if missing:
        errors.append(f"Missing sections: {missing}")

    orig_wc = _word_count(original)
    tail_wc = _word_count(tailored)
    if orig_wc > 0:
        delta_pct = abs(tail_wc - orig_wc) / orig_wc
        if delta_pct > MAX_WORD_DELTA_PCT:
            errors.append(
                f"Word count changed too much: {orig_wc} -> {tail_wc} ({delta_pct:.1%})"
            )

    orig_bullets = _bullet_count(original)
    tail_bullets = _bullet_count(tailored)
    if abs(orig_bullets - tail_bullets) > MAX_BULLET_DELTA:
        errors.append(
            f"Bullet count changed too much: {orig_bullets} -> {tail_bullets}"
        )

    orig_words = set(re.findall(r"\b[A-Z][a-z]{2,}\b", _paragraphs_text(original)))
    tail_words = set(re.findall(r"\b[A-Z][a-z]{2,}\b", _paragraphs_text(tailored)))
    new_words = tail_words - orig_words - {
        "The", "This", "These", "Their", "With", "From", "That", "When",
        "Through", "Building", "Using", "Based", "During", "Within",
    }
    if new_words:
        logger.warning(f"New capitalized words (may be fine): {new_words}")

    return errors


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def tailor_resume(
    company_name: str,
    company_research: dict,
    role_type: str = "security engineering",
) -> str | None:
    """
    Generate a tailored resume DOCX for the given company.
    Returns the output file path on success, None on failure.
    """
    if not BASE_DOCX.exists():
        logger.warning(
            f"Base DOCX not found at {BASE_DOCX}. "
            "Set RESUME_DOCX_PATH env var to your resume .docx path. "
            "Skipping resume tailoring."
        )
        return None

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    original = Document(str(BASE_DOCX))

    skills_idx, skills_text = _extract_skills_paragraph(original)
    summary_idx, summary_text = _extract_summary_paragraph(original)

    logger.info(f"Asking Claude to tailor resume for {company_name}...")
    changes = _call_claude_for_tailoring(
        company_name, company_research, role_type, skills_text, summary_text
    )

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = tmp.name
    shutil.copy(str(BASE_DOCX), tmp_path)

    tailored = Document(tmp_path)
    tailored = _apply_tailoring(tailored, changes)

    errors = _validate(original, tailored)
    if errors:
        logger.error(f"Resume validation failed for {company_name}: {errors}")
        os.unlink(tmp_path)
        return None

    safe_name = re.sub(r"[^\w\s-]", "", company_name).strip().replace(" ", "-")
    out_path = OUTPUT_DIR / f"Francis-Konikkara-{safe_name}-Resume.docx"
    tailored.save(str(out_path))
    os.unlink(tmp_path)

    logger.info(f"Tailored resume saved: {out_path}")
    return str(out_path)
